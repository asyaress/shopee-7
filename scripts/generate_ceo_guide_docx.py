#!/usr/bin/env python3
"""Generate CEO Operations Guide DOCX for Shopee Profit Hub."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "Panduan-Operasional-CEO-Shopee-Profit-Hub.docx"


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_title(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.size = Pt(11)


def add_formula_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, formula) in enumerate(rows):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.text = label
        c1.text = formula
        set_cell_shading(c0, "F4F4F4")
        for cell in (c0, c1):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_flowchart_box(doc, title, steps):
    """Simple vertical flowchart using a styled table."""
    table = doc.add_table(rows=len(steps) * 2 - 1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row_idx = 0
    for i, step in enumerate(steps):
        cell = table.rows[row_idx].cells[0]
        cell.text = step
        set_cell_shading(cell, "F8E4B7" if i == 0 else "FFFFFF")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)
                run.bold = (i == 0)
        row_idx += 1
        if i < len(steps) - 1:
            arrow_cell = table.rows[row_idx].cells[0]
            arrow_cell.text = "▼"
            for p in arrow_cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_idx += 1
    cap = doc.add_paragraph(title)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    doc.add_paragraph()


def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PANDUAN OPERASIONAL CEO\n")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0xAA, 0x3A, 0x3A)
    r2 = title.add_run("Shopee Profit Hub\nEnd-to-End Flow & Rumus Perhitungan")
    r2.font.size = Pt(16)
    r2.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Toedjoe Sinargroup — Dokumen Internal\n").font.size = Pt(11)
    sub.add_run("Versi: Juni 2026 | Aplikasi: Laravel Hub + Android CEO App").font.size = Pt(10)

    doc.add_page_break()

    # TOC
    add_title(doc, "Daftar Isi", 1)
    toc_items = [
        "1. Ringkasan Sistem",
        "2. Flowchart Setup Awal (Sekali)",
        "3. Flowchart Operasional Harian / Mingguan / Bulanan CEO",
        "4. Alur Data Otomatis (Cron & Sync Shopee)",
        "5. Input Manual yang Wajib CEO / Tim",
        "6. Rumus Perhitungan Lengkap",
        "7. Target Bulanan & Pace Tracking",
        "8. HPP, Variant & Packaging",
        "9. Iklan Shopee (Ads) & ROAS",
        "10. Matriks BCG (Performa Produk)",
        "11. Alert CEO Otomatis",
        "12. Checklist Validasi Angka Benar",
        "13. Troubleshooting Umum",
    ]
    add_bullets(doc, toc_items)
    doc.add_page_break()

    # 1. Overview
    add_title(doc, "1. Ringkasan Sistem", 1)
    add_para(doc, (
        "Shopee Profit Hub adalah sistem monitoring laba toko Shopee yang menggabungkan "
        "data order otomatis dari Shopee Open Platform, input biaya manual (HPP, packaging, "
        "operasional), data iklan, dan performa produk (BCG). Output utama: laba kotor/bersih "
        "per toko, per bulan, dan per produk — ditampilkan di web hub dan aplikasi Android CEO."
    ))
    add_para(doc, "Komponen utama:", bold=True)
    add_bullets(doc, [
        "Backend Laravel (shopee.toedjoesinargroup.com) — sync, laporan, kelola data",
        "Database MySQL — orders, products, variants, ads, monthly costs, BCG",
        "Android CEO App — dashboard KPI, HPP priority, target bulanan, alert",
        "Shopee API — Main App (order/produk), Ads App (iklan), AMS (opsional, per produk)",
    ])
    add_para(doc, "Prinsip akurasi angka:", bold=True)
    add_bullets(doc, [
        "Order & fee Shopee = OTOMATIS (dari settlement/financial API)",
        "HPP & packaging = MANUAL (CEO/tim finance) — tanpa ini COGS salah",
        "Biaya operasional bulanan = MANUAL per shop per bulan",
        "Iklan = OTOMATIS sync (butuh permission Ads; per-produk butuh AMS)",
        "BCG akurat = IMPORT Excel Seller Center (auto sync = perkiraan)",
    ])

    # 2. Setup flow
    add_title(doc, "2. Flowchart Setup Awal (Sekali)", 1)
    add_flowchart_box(doc, "Gambar 1 — Setup Awal Sistem", [
        "STEP 0: Deploy server + migrate database",
        "STEP 1: Buat Shopee Open Platform Apps (Main + Ads)",
        "STEP 2: Isi .env (Partner ID, Secret, Redirect, Shop ID)",
        "STEP 3: OAuth / authorize token per toko (Main + Ads)",
        "STEP 4: SHOPEE_CRON_ENABLED=true + cron schedule:run",
        "STEP 5: Backfill data: shopee:sync-all-shops --days=90",
        "STEP 6: Input HPP & packaging semua produk/variant",
        "STEP 7: Input target & operasional bulan berjalan",
        "STEP 8: Install APK CEO + login",
        "STEP 9: Validasi angka vs Seller Center",
    ])

    add_title(doc, "Detail Setup .env Penting", 2)
    env_rows = [
        ("SHOPEE_CRON_ENABLED", "true — aktifkan scheduler sync otomatis"),
        ("SHOPEE_SHOP_ID", "ID toko aktif (contoh: 495488171)"),
        ("SHOPEE_PARTNER_ID / SECRET", "Kredensial Main App"),
        ("SHOPEE_ADS_*", "Kredensial Ads App terpisah"),
        ("SHOPEE_ADS_SYNC_DAYS", "Rentang hari sync iklan (default 30)"),
    ]
    add_formula_table(doc, env_rows)

    add_title(doc, "Authorize Token Shopee", 2)
    add_numbered(doc, [
        "Buka halaman Kelola Data di hub → Connect Shopee (Main App)",
        "Ulangi untuk Ads App jika iklan dipakai",
        "Pastikan token tidak expired — cron gagal jika token invalid",
        "Verifikasi: php artisan shopee:sync-orders --days=1 berhasil tanpa error",
    ])

    doc.add_page_break()

    # 3. CEO operational flow
    add_title(doc, "3. Flowchart Operasional CEO", 1)

    add_title(doc, "3.1 Harian (5–10 menit)", 2)
    add_flowchart_box(doc, "Gambar 2 — Rutinitas Harian CEO", [
        "Buka app CEO → cek KPI hari ini / MTD",
        "Cek alert (bleeder produk, budget ads, HPP missing)",
        "Review order kemarin — ada missing HPP?",
        "Jika produk baru muncul → input HPP variant",
        "Pantau pace target net profit bulan ini",
    ])

    add_title(doc, "3.2 Mingguan (30 menit)", 2)
    add_flowchart_box(doc, "Gambar 3 — Rutinitas Mingguan CEO", [
        "Review ranking produk by net profit",
        "Identifikasi produk rugi (net_profit negatif) → keputusan ads/harga",
        "Cek kelengkasan HPP (target ≥85% complete)",
        "Import BCG Excel dari Seller Center (opsional tapi disarankan)",
        "Review ROAS & ACOS per produk top spend",
    ])

    add_title(doc, "3.3 Bulanan (awal & akhir bulan)", 2)
    add_flowchart_box(doc, "Gambar 4 — Rutinitas Bulanan CEO", [
        "AWAL BULAN: Input shop_monthly_costs (operational + target)",
        "AWAL BULAN: Set target gross, net profit, units, ad budget cap",
        "SEPANJANG BULAN: Pantau progress % vs target",
        "AKHIR BULAN: Rekonsiliasi gross & net vs laporan Shopee",
        "AKHIR BULAN: Update operasional aktual jika ada koreksi",
        "AKHIR BULAN: Export/snapshot keputusan produk (scale/stop ads)",
    ])

    doc.add_page_break()

    # 4. Auto sync
    add_title(doc, "4. Alur Data Otomatis (Cron)", 1)
    add_para(doc, "Scheduler Laravel (app/Console/Kernel.php) — aktif jika SHOPEE_CRON_ENABLED=true:")

    cron_table = doc.add_table(rows=6, cols=3)
    cron_table.style = "Table Grid"
    headers = ["Job", "Frekuensi", "Fungsi"]
    for j, h in enumerate(headers):
        cron_table.rows[0].cells[j].text = h
        set_cell_shading(cron_table.rows[0].cells[j], "A4B787")
    cron_data = [
        ("shopee:sync-all-shops --days=2", "Setiap hari 02:00", "Produk + order + ads semua toko"),
        ("shopee:sync-all-shops --days=7", "Senin 03:00", "Backfill mingguan"),
        ("shopee:sync-orders --days=2", "Setiap 5 menit", "Order incremental"),
        ("shopee:sync-ads --days=2", "Setiap 10 menit", "Data iklan harian"),
        ("ceo:check-alerts", "Setiap hari 08:00", "Alert bleeder/HPP/budget"),
    ]
    for i, row in enumerate(cron_data, start=1):
        for j, val in enumerate(row):
            cron_table.rows[i].cells[j].text = val
    doc.add_paragraph()

    add_flowchart_box(doc, "Gambar 5 — Alur Sync Order → Laba", [
        "Shopee API: get_order_list + get_order_detail",
        "Simpan ke tabel orders + order_items",
        "Ambil financial/settlement → ShopeeOrderFinancial",
        "ShopeeFinancialExtractor → gross, fees, net per order",
        "Match item → product/variant (external_item_id, model_id)",
        "Hitung COGS dari HPP + packaging",
        "Alokasi net ke produk → kurangi ads & operasional",
        "Tampil di Monitoring / CEO App",
    ])

    doc.add_page_break()

    # 5. Manual inputs
    add_title(doc, "5. Input Manual Wajib", 1)
    manual = doc.add_table(rows=7, cols=4)
    manual.style = "Table Grid"
    mh = ["Data", "Siapa", "Kapan", "Di Mana"]
    for j, h in enumerate(mh):
        manual.rows[0].cells[j].text = h
        set_cell_shading(manual.rows[0].cells[j], "AA3A3A")
        for run in manual.rows[0].cells[j].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    manual_data = [
        ("HPP produk", "CEO/Finance", "Saat produk baru / cost berubah", "Web /manage atau App HPP"),
        ("HPP per variant", "CEO/Finance", "Produk multi-SKU", "Web /manage atau App HPP"),
        ("Packaging (fixed/%)", "CEO/Finance", "Saat setup produk", "Web /manage"),
        ("Operasional bulanan", "CEO", "Awal setiap bulan", "shop_monthly_costs / App target"),
        ("Target bulanan", "CEO", "Awal setiap bulan", "target net, gross, units, ad budget"),
        ("BCG Excel import", "Marketing/CEO", "Mingguan", "Seller Center → import hub"),
    ]
    for i, row in enumerate(manual_data, start=1):
        for j, val in enumerate(row):
            manual.rows[i].cells[j].text = val
    doc.add_paragraph()

    add_para(doc, "Urutan input yang BENAR untuk produk baru:", bold=True)
    add_numbered(doc, [
        "Tunggu sync produk dari Shopee (otomatis) — produk muncul di database",
        "CEO input HPP di level variant (jika ada variant) ATAU level produk",
        "CEO input packaging jika ada biaya kemasan terpisah",
        "Tunggu order masuk (sync 5 menit)",
        "Baru angka laba produk tersebut akurat di dashboard",
    ])

    doc.add_page_break()

    # 6. Formulas
    add_title(doc, "6. Rumus Perhitungan Lengkap", 1)
    add_para(doc, "Semua rumus di bawah sesuai kode: ProductProfitReportService, ShopeeFinancialExtractor, MonthlyTargetService.")

    add_title(doc, "6.1 Level Order (Shopee)", 2)
    add_formula_table(doc, [
        ("Gross Produk (gross)", "Total harga barang sebelum potongan Shopee (dari financial API)"),
        ("Gross Buyer", "Total yang dibayar pembeli (termasuk ongkir buyer jika ada)"),
        ("Fee Total", "fee_admin + fee_program_hemat + fee_service + fee_process + fee_ams + fee_campaign + fee_premi + fee_seller_transaction"),
        ("Net Order (net)", "Pendapatan bersih seller setelah fee Shopee (dari API escrow_amount / seller income)"),
        ("Catatan non-Shopee", "Order jenis non-shopee: fee = 0, net = gross"),
    ])

    add_title(doc, "6.2 COGS (HPP + Packaging) per Line Item", 2)
    add_formula_table(doc, [
        ("Unit HPP", "variant.hpp_amount ?? product.hpp_amount (variant menang jika ada)"),
        ("Unit Packaging (fixed)", "packaging_value langsung per unit"),
        ("Unit Packaging (percent)", "unit_price × (packaging_value / 100)"),
        ("Line COGS", "(unit_hpp + unit_packaging) × quantity"),
        ("Missing cost flag", "TRUE jika produk tidak ketemu ATAU hpp & packaging keduanya null"),
    ])
    add_para(doc, "Contoh: HPP Rp 8.000, packaging fixed Rp 500, qty 3 → COGS = (8000+500)×3 = Rp 25.500")

    add_title(doc, "6.3 Alokasi Net ke Produk", 2)
    add_para(doc, (
        "Net order dialokasikan ke setiap produk berdasarkan proporsi gross line item "
        "terhadap total gross item di order yang sama:"
    ))
    add_formula_table(doc, [
        ("Ratio produk", "gross_line_item / sum(gross_all_items_in_order)"),
        ("Net alokasi produk", "net_order × ratio"),
        ("Gross profit produk", "net_alokasi − COGS_produk"),
    ])
    add_para(doc, "Item tidak ter-match produk → masuk bucket 'unknown' dengan alokasi serupa.")

    add_title(doc, "6.4 Laba Kotor & Bersih (Shop Level)", 2)
    add_formula_table(doc, [
        ("Gross profit shop", "Σ net_order − Σ COGS_order  (= Σ gross_profit per order)"),
        ("Ads total", "Σ spend dari shopee_product_ads_daily dalam rentang tanggal"),
        ("Operasional total", "Σ operational_amount dari shop_monthly_costs untuk setiap bulan yang tersentuh rentang (FULL bulan, tidak di-prorate harian)"),
        ("Net profit shop", "gross_profit_shop − ads_total − operational_total"),
        ("Take rate", "fee_total / gross"),
        ("Gross margin %", "gross_profit / gross"),
        ("Net margin %", "net_profit / gross"),
    ])

    add_title(doc, "6.5 Laba per Produk (setelah ads & ops)", 2)
    add_formula_table(doc, [
        ("Ads produk", "SUM(spend) per product_id / external_item_id dari shopee_product_ads_daily"),
        ("Operasional alokasi", "operational_total × (net_produk / net_shop_total)"),
        ("Net profit produk", "gross_profit_produk − ads_produk − operasional_alokasi"),
        ("Margin produk", "net_profit_produk / net_produk (jika net > 0)"),
        ("ROAS", "gross_produk / ads_spend (jika ads > 0)"),
        ("ACOS", "ads_spend / gross_produk (jika gross > 0)"),
    ])

    add_title(doc, "6.6 Metrik Bulanan (Pivot)", 2)
    add_formula_table(doc, [
        ("Net profit bulan", "gross_profit_bulan − ads_bulan − operational_bulan"),
        ("AOV gross", "gross_bulan / jumlah_order_bulan"),
        ("Basket size", "units_bulan / jumlah_order_bulan"),
        ("Gross margin % bulan", "gross_profit_bulan / gross_bulan"),
        ("Net margin % bulan", "net_profit_bulan / gross_bulan"),
    ])

    doc.add_page_break()

    # 7. Monthly targets
    add_title(doc, "7. Target Bulanan & Pace Tracking", 1)
    add_para(doc, "Data disimpan di tabel shop_monthly_costs per shop_id + year_month (format YYYY-MM).")
    add_formula_table(doc, [
        ("Progress net %", "actual_net_profit / target_net_profit"),
        ("Progress gross %", "actual_gross / target_gross"),
        ("Progress units %", "actual_units / target_units"),
        ("Progress ads %", "actual_ads / ad_budget_cap"),
        ("Pace factor", "hari_ini / jumlah_hari_dalam_bulan"),
        ("Expected net by today", "target_net_profit × pace_factor"),
        ("On track?", "actual_net ≥ expected_net × 0.9 ( toleransi 10% di bawah pace )"),
    ])
    add_para(doc, "Actual dihitung dari awal bulan s/d hari ini (bukan full month projection).")

    # 8. HPP
    add_title(doc, "8. HPP, Variant & Packaging", 1)
    add_formula_table(doc, [
        ("Produk complete?", "Jika punya variant: minimal 1 variant punya hpp_amount ≥ 0. Jika tidak: product.hpp_amount tidak null"),
        ("Gate rekomendasi", "≥85% produk complete → gate_ok. ≥70% → recommendations_allowed"),
        ("Priority list", "Produk missing HPP dengan penjualan tertinggi ditampilkan di app CEO"),
        ("Variant override", "HPP/packaging variant menimpa level produk induk"),
    ])
    add_flowchart_box(doc, "Gambar 6 — Flow Input HPP Produk Baru", [
        "Sync produk dari Shopee",
        "Cek: produk punya variant?",
        "YA → input HPP tiap variant yang dijual",
        "TIDAK → input HPP di level produk",
        "Input packaging jika perlu",
        "Refresh dashboard → missing_cost = false",
    ])

    doc.add_page_break()

    # 9. Ads
    add_title(doc, "9. Iklan Shopee (Ads)", 1)
    add_para(doc, (
        "Sync iklan via Ads API. Data disimpan di shopee_product_ads_daily. "
        "Spend dijumlahkan per hari per produk (atau shop_aggregate jika per-produk tidak tersedia)."
    ))
    add_bullets(doc, [
        "ROAS tinggi = gross produk besar relatif terhadap spend iklan",
        "ACOS rendah = spend iklan kecil relatif terhadap gross",
        "Produk 'bleeder' = net_profit negatif setelah ads + operasional",
        "Permission AMS (get_product_performance): diperlukan untuk data ads native per item",
        "Tanpa AMS: hanya shop_aggregate atau estimasi campaign split — label sebagai perkiraan",
    ])

    # 10. BCG
    add_title(doc, "10. Matriks BCG", 1)
    add_formula_table(doc, [
        ("Conversion rate", "dari import Excel ATAU units_sold / visitors jika rate = 0"),
        ("Traffic baseline", "Median visitors produk aktif (config monitoring.bcg_funnel)"),
        ("Threshold konversi", "Default 2% (config conversion_threshold)"),
        ("STAR", "Konversi tinggi + traffic tinggi"),
        ("CASH COW", "Konversi tinggi + traffic rendah"),
        ("QUESTION MARK", "Konversi rendah + traffic tinggi"),
        ("DOG", "Konversi rendah + traffic rendah"),
    ])
    add_para(doc, "Sumber data: SOURCE_IMPORT (akurat dari Seller Center) vs SOURCE_AUTO (sync API — perkiraan).")

    # 11. Alerts
    add_title(doc, "11. Alert CEO Otomatis", 1)
    add_para(doc, "Command ceo:check-alerts jalan setiap hari 08:00. Cek CeoAlertService untuk kondisi:")
    add_bullets(doc, [
        "Produk bleeder — net profit negatif dengan spend ads signifikan",
        "Budget ads mendekati/melebihi ad_budget_cap bulan ini",
        "HPP completeness di bawah threshold gate",
        "Target net profit off-track vs pace",
    ])

    doc.add_page_break()

    # 12. Validation checklist
    add_title(doc, "12. Checklist Validasi — Angka Harus Benar", 1)
    add_para(doc, "CEO lakukan checklist ini sebelum percaya dashboard 100%:", bold=True)
    checklist = [
        "[ ] Token Shopee Main + Ads valid (sync manual sukses)",
        "[ ] Cron schedule:run aktif di server (crontab)",
        "[ ] Order 7 hari terakhir count ≈ Seller Center",
        "[ ] HPP completeness ≥ 85%",
        "[ ] Operasional bulan berjalan sudah diinput",
        "[ ] Target bulan berjalan sudah diinput",
        "[ ] Gross MTD hub vs Seller Center selisih < 5% (fee timing)",
        "[ ] Produk top 10 punya HPP + ads data (jika running ads)",
        "[ ] BCG diimport mingguan (jika pakai keputusan funnel)",
        "[ ] App CEO login ke shop_id benar",
    ]
    add_bullets(doc, checklist)

    add_title(doc, "Urutan Validasi End-to-End", 2)
    add_flowchart_box(doc, "Gambar 7 — Validasi End-to-End", [
        "1. Sync order → bandingkan 1 order sample fee breakdown",
        "2. Input HPP sample → COGS order match manual",
        "3. Input operasional → net profit shop turun sesuai nominal",
        "4. Sync ads → ads total muncul di summary",
        "5. Cek 1 produk: net alokasi + COGS + ads + ops = net_profit",
        "6. Bandingkan MTD net profit vs target pace",
        "SEMUA OK → sistem siap operasional",
    ])

    # 13. Troubleshooting
    add_title(doc, "13. Troubleshooting", 1)
    trouble = doc.add_table(rows=8, cols=2)
    trouble.style = "Table Grid"
    trouble.rows[0].cells[0].text = "Gejala"
    trouble.rows[0].cells[1].text = "Solusi"
    set_cell_shading(trouble.rows[0].cells[0], "F4F4F4")
    set_cell_shading(trouble.rows[0].cells[1], "F4F4F4")
    fixes = [
        ("Order tidak masuk", "Cek token, SHOPEE_CRON_ENABLED, php artisan schedule:list, jalankan shopee:sync-orders manual"),
        ("Variant tidak muncul di app", "git pull di server, php artisan config:clear, pastikan API ceo/hpp/priority include variants"),
        ("Missing cost di banyak produk", "Lengkapi HPP via /manage atau app CEO HPP priority"),
        ("Ads hanya shop_aggregate", "Minta permission AMS ke Shopee atau terima estimasi campaign-level"),
        ("Net profit terlalu tinggi", "Belum input operasional bulan ini"),
        ("Net profit terlalu rendah", "Cek HPP terlalu tinggi, double count operasional multi-bulan di rentang filter"),
        ("Angka beda Seller Center", "Timing settlement, filter status order, rentang tanggal order_date vs payout"),
    ]
    for i, (sym, fix) in enumerate(fixes, start=1):
        trouble.rows[i].cells[0].text = sym
        trouble.rows[i].cells[1].text = fix

    doc.add_paragraph()
    add_para(doc, "Command berguna untuk debugging:", bold=True)
    add_bullets(doc, [
        "php artisan shopee:sync-orders --days=7",
        "php artisan shopee:sync-ads --days=7",
        "php artisan shopee:debug-ads {shop_id}",
        "php artisan schedule:list",
        "php artisan ceo:check-alerts",
    ])

    # Footer note
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("— Akhir Dokumen —").italic = True
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(str(OUTPUT))
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
